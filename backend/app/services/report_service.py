"""Service for generating reports and PDFs."""
import os
import re
import time
from datetime import datetime
from pathlib import Path
from jinja2 import Template, Environment, BaseLoader
from html.parser import HTMLParser
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None
from sqlalchemy import and_, or_
from app import db
from app.models import Report, ReportTemplate, Timeline, TimelineEntry, Project, User
from app.models.mitre import TimelineAnalysisResult


class ReportService:
    """Service for generating formatted reports and PDFs."""
    
    def __init__(self, reports_dir=None):
        """Initialize the report service.
        
        Args:
            reports_dir: Directory to store generated PDF files. Defaults to 'reports' in app root.
        """
        if reports_dir is None:
            reports_dir = Path(__file__).parent.parent.parent / 'reports'
        
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(exist_ok=True)
    
    @staticmethod
    def validate_llm_report_data(project_id, timeline_id=None):
        """Validate that there is sufficient analyzed data for LLM report generation.
        
        Args:
            project_id: ID of the project
            timeline_id: Optional specific timeline ID
            
        Returns:
            Dict with 'valid' (bool), 'message' (str), and optional 'details' (dict)
        """
        # Build query for significant event IDs
        # We look for:
        # 1. High/Critical priority (score >= 0.6)
        # 2. MITRE ATT&CK mappings
        
        stmt = db.session.query(TimelineAnalysisResult.entry_id).join(TimelineEntry)
        
        if timeline_id:
            stmt = stmt.filter(TimelineEntry.timeline_id == timeline_id)
        else:
            stmt = stmt.join(Timeline).filter(Timeline.project_id == project_id)
            
        stmt = stmt.filter(
            db.or_(
                db.and_(
                    TimelineAnalysisResult.analysis_type == 'prioritization',
                    TimelineAnalysisResult.priority_score >= 0.6
                ),
                db.and_(
                    TimelineAnalysisResult.analysis_type == 'attack_mapping',
                    TimelineAnalysisResult.mitre_technique_id.isnot(None)
                )
            )
        ).distinct()
        
        count = stmt.count()
        
        if count == 0:
            scope = f"timeline {timeline_id}" if timeline_id else f"project {project_id}"
            return {
                'valid': False,
                'message': f"No analyzed events found for {scope}. Please run event prioritization or MITRE ATT&CK mapping before generating an AI report.",
                'details': {
                    'analyzed_events': 0,
                    'timeline_id': timeline_id,
                    'project_id': project_id
                }
            }
        
        return {
            'valid': True,
            'message': 'Sufficient analyzed data available',
            'details': {
                'analyzed_events': count
            }
        }
        
    def generate_report(self, template_id, project_id, user_id, parameters):
        """Generate a report from a template.
        
        Args:
            template_id: ID of the report template to use
            project_id: ID of the project
            user_id: ID of the user generating the report
            parameters: Dict with report parameters:
                - timeline_id (optional): Specific timeline to report on
                - start_date (optional): Filter entries after this date
                - end_date (optional): Filter entries before this date
                - entry_limit (optional): Max number of entries to include
                - filters (optional): Additional filters for entries
                - name: Name for the generated report
                - description (optional): Description of the report
                - format (optional): Output format ('pdf' or 'docx', defaults to 'pdf')
        
        Returns:
            Report: The generated report instance
        """
        start_time = time.time()
        
        # Load template
        template = ReportTemplate.query.get(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        # Load project
        project = Project.query.get(project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")
        
        # Fetch data based on parameters
        data = self._fetch_report_data(project_id, parameters)
        
        # Render HTML content
        html_content = self._render_template(template, data, parameters)
        
        # Determine output format
        output_format = parameters.get('format', 'pdf').lower()
        if output_format not in ['pdf', 'docx']:
            output_format = 'pdf'
        
        # Create report record
        report = Report(
            name=parameters.get('name', f"Report - {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
            description=parameters.get('description'),
            template_id=template_id,
            project_id=project_id,
            timeline_id=parameters.get('timeline_id'),
            created_by=user_id,
            parameters=parameters,
            html_content=html_content,
            entry_count=len(data['entries']),
            format=output_format
        )
        
        db.session.add(report)
        db.session.flush()  # Get the report ID
        
        # Generate file based on format
        try:
            if output_format == 'docx':
                file_path = self._generate_docx(html_content, data, report.id, template.config or {})
            else:
                file_path = self._generate_pdf(html_content, report.id, template.config or {})
        except Exception as e:
            # If generation fails, raise with clear message
            db.session.rollback()
            raise ValueError(f"Failed to generate {output_format.upper()} report: {str(e)}")
        
        report.file_path = file_path
        report.pdf_path = file_path if output_format == 'pdf' else None  # Backwards compatibility
        
        # Get file size
        full_path = self.reports_dir / file_path
        if full_path.exists():
            report.file_size = full_path.stat().st_size
        
        report.generation_time = time.time() - start_time
        
        db.session.commit()
        
        return report
    
    async def generate_llm_report(self, project_id, user_id, parameters):
        """Generate an AI-powered report.
        
        Args:
            project_id: ID of the project
            user_id: ID of the user generating the report
            parameters: Dict with report parameters
            
        Returns:
            Report: The generated report instance
        """
        from app.services.timeline_analysis_service import TimelineAnalysisService
        
        start_time = time.time()
        timeline_service = TimelineAnalysisService()
        
        # Generate content
        content = await timeline_service.generate_case_report_content(
            project_id=project_id,
            timeline_id=parameters.get('timeline_id'),
            context=parameters.get('description')
        )
        
        # Render HTML
        html_content = self._render_llm_html(content, parameters)
        
        # Create report record
        report = Report(
            name=parameters.get('name', f"AI Investigation Report - {datetime.now().strftime('%Y-%m-%d')}"),
            description=parameters.get('description', 'AI-generated forensic report'),
            project_id=project_id,
            timeline_id=parameters.get('timeline_id'),
            created_by=user_id,
            parameters=parameters,
            html_content=html_content,
            format='docx',  # Default to DOCX for AI reports
            entry_count=0  # TODO: Update if we track which events were used
        )
        
        db.session.add(report)
        db.session.flush()
        
        # Generate DOCX
        # Pass empty data dict as we don't use the standard data structure for this template
        file_path = self._generate_docx(html_content, {}, report.id, {})
        report.file_path = file_path
        
        # Size
        full_path = self.reports_dir / file_path
        if full_path.exists():
            report.file_size = full_path.stat().st_size
            
        report.generation_time = time.time() - start_time
        db.session.commit()
        
        return report

    def _render_llm_html(self, content, parameters):
        """Render the LLM content into HTML."""
        template = """
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; line-height: 1.6; }
                h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
                h2 { color: #34495e; margin-top: 20px; }
                .meta { color: #7f8c8d; margin-bottom: 30px; }
                .severity-High { color: #e74c3c; font-weight: bold; }
                .severity-Medium { color: #f39c12; font-weight: bold; }
                .severity-Low { color: #27ae60; font-weight: bold; }
                .finding { margin-bottom: 20px; border-bottom: 1px solid #eee; padding-bottom: 20px; }
                ul { margin-top: 10px; }
                li { margin-bottom: 5px; }
            </style>
        </head>
        <body>
            <h1>Investigation Report</h1>
            <div class="meta">
                <p><strong>Generated:</strong> {{ date }}</p>
                <p><strong>Case Context:</strong> {{ context }}</p>
            </div>

            <h1>Executive Summary</h1>
            <div>{{ executive_summary }}</div>

            <h1>Key Findings</h1>
            {% for finding in key_findings %}
            <div class="finding">
                <h2>{{ finding.title }}</h2>
                <p><strong>MITRE Phase:</strong> {{ finding.mitre_phase }}</p>
                <p><strong>Severity:</strong> <b class="severity-{{ finding.severity }}">{{ finding.severity }}</b></p>
                <div>{{ finding.description }}</div>
            </div>
            {% endfor %}

            <h1>Recommendations</h1>
            <ul>
            {% for rec in recommendations %}
                <li>{{ rec }}</li>
            {% endfor %}
            </ul>
        </body>
        </html>
        """
        
        env = Environment(loader=BaseLoader())
        jinja_template = env.from_string(template)
        
        return jinja_template.render(
            date=datetime.now().strftime('%Y-%m-%d %H:%M'),
            context=parameters.get('description', 'N/A'),
            executive_summary=content.get('executive_summary', '').replace('\n', '<br>'),
            key_findings=content.get('key_findings', []),
            recommendations=content.get('recommendations', [])
        )
    
    def _fetch_report_data(self, project_id, parameters):
        """Fetch timeline entries and related data for the report.
        
        Args:
            project_id: Project ID
            parameters: Report parameters with filters
            
        Returns:
            Dict with entries, project info, timelines, stats, etc.
        """
        # Base query for entries
        query = db.session.query(TimelineEntry).join(Timeline).filter(
            Timeline.project_id == project_id
        )
        
        # Filter by timeline if specified
        if parameters.get('timeline_id'):
            query = query.filter(TimelineEntry.timeline_id == parameters['timeline_id'])
        
        # Date range filters
        if parameters.get('start_date'):
            start_date = datetime.fromisoformat(parameters['start_date'].replace('Z', '+00:00'))
            query = query.filter(
                or_(
                    TimelineEntry.data['Timestamp'].astext >= start_date.isoformat(),
                    TimelineEntry.data['TimeCreated'].astext >= start_date.isoformat()
                )
            )
        
        if parameters.get('end_date'):
            end_date = datetime.fromisoformat(parameters['end_date'].replace('Z', '+00:00'))
            query = query.filter(
                or_(
                    TimelineEntry.data['Timestamp'].astext <= end_date.isoformat(),
                    TimelineEntry.data['TimeCreated'].astext <= end_date.isoformat()
                )
            )
        
        # Apply entry limit
        entry_limit = parameters.get('entry_limit', 1000)
        query = query.order_by(TimelineEntry.id.desc()).limit(entry_limit)
        
        entries = query.all()
        
        # Get project and timelines
        project = Project.query.get(project_id)
        timelines = Timeline.query.filter_by(project_id=project_id).all()
        
        # Calculate statistics
        stats = self._calculate_stats(entries, timelines)
        
        return {
            'entries': [self._entry_to_dict(e) for e in entries],
            'project': {
                'id': project.id,
                'name': project.name,
                'description': project.description
            },
            'timelines': [{'id': t.id, 'name': t.name} for t in timelines],
            'stats': stats,
            'generated_at': datetime.now().isoformat(),
            'parameters': parameters
        }
    
    def _entry_to_dict(self, entry):
        """Convert a TimelineEntry to a dict for template rendering."""
        data = entry.data.copy() if entry.data else {}
        return {
            'id': entry.id,
            'timeline_id': entry.timeline_id,
            'data': data,
            'created_at': entry.created_at.isoformat() if entry.created_at else None
        }
    
    def _calculate_stats(self, entries, timelines):
        """Calculate statistics for the report."""
        timeline_counts = {}
        for entry in entries:
            timeline_id = entry.timeline_id
            timeline_counts[timeline_id] = timeline_counts.get(timeline_id, 0) + 1
        
        return {
            'total_entries': len(entries),
            'total_timelines': len(timelines),
            'entries_by_timeline': timeline_counts
        }
    
    def _render_template(self, template, data, parameters):
        """Render the template with data using Jinja2.
        
        Args:
            template: ReportTemplate instance
            data: Data dict to pass to template
            parameters: Report parameters
            
        Returns:
            Rendered HTML string
        """
        env = Environment(loader=BaseLoader())
        jinja_template = env.from_string(template.template_content)
        
        # Render with data
        html = jinja_template.render(
            data=data,
            project=data['project'],
            entries=data['entries'],
            timelines=data['timelines'],
            stats=data['stats'],
            parameters=parameters,
            generated_at=data['generated_at']
        )
        
        return html
    
    def _generate_pdf(self, html_content, report_id, config):
        """Generate PDF from HTML content.
        
        Args:
            html_content: Rendered HTML string
            report_id: ID of the report
            config: Template configuration with page settings
            
        Returns:
            Relative path to the generated PDF file
        """
        # Generate filename
        filename = f"report_{report_id}_{int(time.time())}.pdf"
        filepath = self.reports_dir / filename
        
        # Get configuration
        page_size = config.get('page_size', 'A4')
        orientation = config.get('orientation', 'portrait')
        
        # Add basic styling if not present in HTML
        if '<style>' not in html_content and '<link' not in html_content:
            html_content = self._add_default_styles(html_content)
        
        # Generate PDF
        html = HTML(string=html_content)
        
        # CSS for page setup
        css_string = f"""
        @page {{
            size: {page_size} {orientation};
            margin: 2cm;
        }}
        """
        css = CSS(string=css_string)
        
        html.write_pdf(str(filepath), stylesheets=[css])
        
        return filename
    
    def _add_default_styles(self, html_content):
        """Add default CSS styles to HTML content."""
        default_css = """
        <style>
            body {
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
            }
            h1 {
                color: #2c3e50;
                border-bottom: 2px solid #3498db;
                padding-bottom: 10px;
            }
            h2 {
                color: #34495e;
                margin-top: 20px;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            th {
                background-color: #3498db;
                color: white;
            }
            tr:nth-child(even) {
                background-color: #f2f2f2;
            }
            .stats {
                background-color: #ecf0f1;
                padding: 15px;
                border-radius: 5px;
                margin: 20px 0;
            }
            .timestamp {
                color: #7f8c8d;
                font-size: 0.9em;
            }
        </style>
        """
        
        # Insert styles after <head> or at the beginning
        if '<head>' in html_content:
            html_content = html_content.replace('<head>', f'<head>{default_css}')
        elif '<body>' in html_content:
            html_content = html_content.replace('<body>', f'{default_css}<body>')
        else:
            html_content = default_css + html_content
        
        return html_content
    
    def get_pdf_path(self, report_id):
        """Get the full filesystem path for a report file (PDF or DOCX).
        
        Args:
            report_id: ID of the report
            
        Returns:
            Path object to the report file, or None if not found
        """
        report = Report.query.get(report_id)
        if not report:
            return None
        
        # Use file_path first (supports both formats), fallback to pdf_path for backwards compatibility
        file_path = report.file_path or report.pdf_path
        if not file_path:
            return None
        
        full_path = self.reports_dir / file_path
        return full_path if full_path.exists() else None
    
    def delete_report_pdf(self, report_id):
        """Delete the PDF file for a report.
        
        Args:
            report_id: ID of the report
            
        Returns:
            True if deleted, False if not found
        """
        pdf_path = self.get_pdf_path(report_id)
        if pdf_path and pdf_path.exists():
            pdf_path.unlink()
            return True
        return False
    
    def _generate_docx(self, html_content, data, report_id, config):
        """Generate DOCX from HTML content and data.
        
        Args:
            html_content: Rendered HTML string
            data: Report data dict
            report_id: ID of the report
            config: Template configuration
            
        Returns:
            Relative path to the generated DOCX file
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed")
        
        # Generate filename
        filename = f"report_{report_id}_{int(time.time())}.docx"
        filepath = self.reports_dir / filename
        
        # Create document
        document = Document()
        
        # Parse HTML and convert to Word document
        parser = HTMLToDocxParser(document, data)
        parser.feed(html_content)
        
        # Save document
        document.save(str(filepath))
        
        return filename


class HTMLToDocxParser(HTMLParser):
    """Parse HTML and convert to Word document format."""
    
    def __init__(self, document, data):
        super().__init__()
        self.document = document
        self.data = data
        self.current_paragraph = None
        self.current_run = None
        self.in_heading = None
        self.in_table = False
        self.table = None
        self.table_row = None
        self.table_cell = None
        self.table_rows = []  # Buffer to collect rows before creating table
        self.current_table_row = []  # Current row being built
        self.list_level = 0
        self.bold = False
        self.italic = False
        self.underline = False
        self.ignore_content = False
    
    def handle_starttag(self, tag, attrs):
        """Handle opening HTML tags."""
        if tag in ['style', 'script', 'head', 'meta', 'title']:
            self.ignore_content = True
            return
            
        if self.ignore_content:
            return

        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_heading = int(tag[1])
            self.current_paragraph = self.document.add_heading('', level=self.in_heading)
        elif tag == 'p':
            self.current_paragraph = self.document.add_paragraph()
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
            else:
                self.document.add_paragraph()
        elif tag == 'b' or tag == 'strong':
            self.bold = True
        elif tag == 'i' or tag == 'em':
            self.italic = True
        elif tag == 'u':
            self.underline = True
        elif tag == 'table':
            self.in_table = True
            self.table_rows = []
        elif tag == 'tr':
            self.current_table_row = []
        elif tag in ['th', 'td']:
            # Start collecting cell content
            self.table_cell_content = ''
            if tag == 'th':
                self.bold = True
        elif tag in ['ul', 'ol']:
            self.list_level += 1
        elif tag == 'li':
            self.current_paragraph = self.document.add_paragraph(style='List Bullet' if self.list_level > 0 else 'Normal')
    
    def handle_endtag(self, tag):
        """Handle closing HTML tags."""
        if tag in ['style', 'script', 'head', 'meta', 'title']:
            self.ignore_content = False
            return

        if self.ignore_content:
            return

        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_heading = None
            self.current_paragraph = None
        elif tag == 'p':
            self.current_paragraph = None
        elif tag == 'b' or tag == 'strong':
            self.bold = False
        elif tag == 'i' or tag == 'em':
            self.italic = False
        elif tag == 'u':
            self.underline = False
        elif tag == 'table':
            # Create the table now that we know dimensions
            if self.table_rows:
                max_cols = max(len(row) for row in self.table_rows)
                table = self.document.add_table(rows=len(self.table_rows), cols=max_cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(self.table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text
            
            self.in_table = False
            self.table_rows = []
        elif tag == 'tr':
            if self.current_table_row is not None:
                self.table_rows.append(self.current_table_row)
            self.current_table_row = None
        elif tag in ['th', 'td']:
            if hasattr(self, 'table_cell_content'):
                self.current_table_row.append(self.table_cell_content)
                delattr(self, 'table_cell_content')
            if tag == 'th':
                self.bold = False
        elif tag in ['ul', 'ol']:
            self.list_level -= 1
        elif tag == 'li':
            self.current_paragraph = None
    
    def handle_data(self, data):
        """Handle text data."""
        if self.ignore_content:
            return

        data = data.strip()
        if not data:
            return
        
        # If we're in a table cell, collect the content
        if hasattr(self, 'table_cell_content'):
            self.table_cell_content += data
            return
        
        if not self.current_paragraph:
            self.current_paragraph = self.document.add_paragraph()
        
        run = self.current_paragraph.add_run(data)
        
        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True
        if self.underline:
            run.underline = True
